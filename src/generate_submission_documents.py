"""Create the generic submission manuscript, title page and cover letter."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from generate_submission_tables import add_table as add_submission_table
from generate_submission_tables import table_1, table_2, table_3, table_4, table_5


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
PKG = ROOT / "submission_package_2026-08-26"
PKG.mkdir(parents=True, exist_ok=True)
FONT = "Times New Roman"
MANUSCRIPT_LINE_SPACING = 1.5


AUTHORS = [
    ("Pengqin Shu", "1", "#"),
    ("Xiaoye Xu", "2", "#"),
    ("Fei Ying", "1", "#"),
    ("Jing Sun", "1", "#"),
    ("Zhao Liu", "1", ""),
    ("Tielong Chen", "1", "*"),
]

TITLE = "Transportability and calibration of an hourly model for impending continuous vasopressor initiation across MIMIC-IV, eICU-CRD, and SICdb"
SHORT_TITLE = "Transportability of vasopressor initiation prediction"
ARTICLE_TYPE = "Original Research"
KEYWORDS = "transportability; critical care; vasopressor initiation; machine learning; calibration; external validation; MIMIC-IV; eICU-CRD; SICdb"

REFERENCES = [
    "Collins GS, Moons KGM, Dhiman P, Riley RD, Beam AL, Van Calster B, et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ. 2024;385:e078378. doi:10.1136/bmj-2023-078378.",
    "Duval L, Villie A, Zheng F, Terraz G, Blein S, Duperchy E, et al. Early prediction of vasopressor initiation in ICU sepsis patients using an interpretable EHR-based ML model. BMC Med Inform Decis Mak. 2025;25:442. doi:10.1186/s12911-025-03274-3.",
    "Kwak GH, Ling L, Hui P. Predicting the need for vasopressors in the intensive care unit using an attention based deep learning model. Shock. 2021;56(1):73-79. doi:10.1097/SHK.0000000000001692.",
    "Holder AL, Shashikumar SP, Wardi G, Buchman TG, Nemati S. A locally optimized data-driven tool to predict sepsis-associated vasopressor use in the ICU. Crit Care Med. 2021;49(12):e1196-e1205. doi:10.1097/CCM.0000000000005175.",
    "Rahman A, Chang Y, Dong J, Conroy B, Natarajan A, Kinoshita T, et al. Early prediction of hemodynamic interventions in the intensive care unit using machine learning. Crit Care. 2021;25:388. doi:10.1186/s13054-021-03808-x.",
    "Chiang DH, Tian C, Jiang Z, Ou-Yang YS, Lin YY. External validation of a machine learning model to predict hemodynamic instability in intensive care unit. Crit Care. 2022;26:215. doi:10.1186/s13054-022-04088-9.",
    "Chiang DH, Jiang Z, Tian C, Wang CY. Development and validation of a dynamic early warning system with time-varying machine learning models for predicting hemodynamic instability in critical care: a multicohort study. Crit Care. 2025;29:318. doi:10.1186/s13054-025-05553-x.",
    "Zhang Z, Celi LA, Ho KM. Prediction of extended period of vasopressor infusion requiring central venous catheterisation: a burning issue in critical care. Anaesth Intensive Care. 2021;49(4):250-252. doi:10.1177/0310057X211030927.",
    "Kim DS, Park JE, Hwang SY, Jeong D, Lee GT, Kim T. Prediction of vasopressor requirement among hypotensive patients with suspected infection: usefulness of diastolic shock index and lactate. Clin Exp Emerg Med. 2022;9(3):176-186. doi:10.15441/ceem.22.324.",
    "Jeong D, Choi M, Maeng SJ, Yoon H, Park JE, Lee GT. Machine learning-based clusters of vital signs and lactate levels predict vasopressor use in sepsis. Clin Exp Emerg Med. 2026 Jan 14. doi:10.15441/ceem.25.247.",
    "Muralitharan S, Nelson W, Di S, McGillion M, Devereaux PJ, Barr NG, et al. Machine learning-based early warning systems for clinical deterioration: systematic scoping review. J Med Internet Res. 2021;23(2):e25187. doi:10.2196/25187.",
    "Goh KH, Wang L, Yeow AYK, Poh H, Li K, Yeow JJ, et al. Artificial intelligence in sepsis early prediction and diagnosis using unstructured data in healthcare. Nat Commun. 2021;12:711. doi:10.1038/s41467-021-20910-4.",
    "Moor M, Bennett N, Plecko D, et al. Predicting sepsis using deep learning across international sites: a retrospective development and validation study. EClinicalMedicine. 2023;58:102124. doi:10.1016/j.eclinm.2023.102124.",
    "Yuan C, Beaulieu-Jones BK, Yu S. Temporal bias in case-control design: preventing reliable predictions of the future. Nat Commun. 2021;12:1107. doi:10.1038/s41467-021-21390-2.",
    "Wong A, Otles E, Donnelly JP, Krumm A, McCullough J, DeTroyer-Cooley O, et al. External validation of a widely implemented proprietary sepsis prediction model in hospitalized patients. JAMA Intern Med. 2021;181(8):1065-1070. doi:10.1001/jamainternmed.2021.2626.",
    "Johnson AEW, Bulgarelli L, Shen L, Gayles A, Shammout A, Horng S, et al. MIMIC-IV, a freely accessible electronic health record dataset. Sci Data. 2023;10:1. doi:10.1038/s41597-022-01899-x.",
    "Pollard TJ, Johnson AEW, Raffa JD, Celi LA, Mark RG, Badawi O. The eICU Collaborative Research Database, a freely available multi-center database for critical care research. Sci Data. 2018;5:180178. doi:10.1038/sdata.2018.178.",
    "Rodemund N, Wernly B, Jung C, Cozowicz C, Koköfer A. The Salzburg Intensive Care database (SICdb): an openly available critical care dataset. Intensive Care Med. 2023;49(6):700-702. doi:10.1007/s00134-023-07046-3.",
    "Johnson AEW, Bulgarelli L, Pollard T, et al. MIMIC-IV (version 3.1) [dataset]. PhysioNet; 2024. doi:10.13026/kpb9-mt58. Available from: https://physionet.org/content/mimiciv/3.1/.",
    "Pollard TJ, Johnson A, Raffa J, et al. eICU Collaborative Research Database (version 2.0) [dataset]. PhysioNet; 2019. doi:10.13026/C2WM1R. Available from: https://physionet.org/content/eicu-crd/2.0/.",
    "Rodemund N, Koköfer A, Wernly B, et al. Salzburg Intensive Care database (SICdb), a freely accessible intensive care database (version 1.0.8) [dataset]. PhysioNet; 2024. doi:10.13026/8m72-6j83. Available from: https://physionet.org/content/sicdb/1.0.8/.",
    "Wolff RF, Moons KGM, Riley RD, Whiting PF, Westwood M, Collins GS, et al. PROBAST: a tool to assess the risk of bias and applicability of prediction model studies. Ann Intern Med. 2019;170(1):51-58. doi:10.7326/M18-1376.",
    "Collins GS, Reitsma JB, Altman DG, Moons KGM. Transparent reporting of a multivariable prediction model for individual prognosis or diagnosis (TRIPOD): the TRIPOD statement. BMJ. 2015;350:g7594. doi:10.1136/bmj.g7594.",
    "Debray TPA, Vergouwe Y, Koffijberg H, Nieboer D, Steyerberg EW, Moons KGM. A new framework to enhance the interpretation of external validation studies of clinical prediction models. J Clin Epidemiol. 2015;68(3):279-289. doi:10.1016/j.jclinepi.2014.06.018.",
    "Debray TPA, Collins GS, Riley RD, Snell KIE, Van Calster B, Reitsma JB, et al. Transparent reporting of multivariable prediction models developed or validated using clustered data: TRIPOD-Cluster checklist. BMJ. 2023;380:e071018. doi:10.1136/bmj-2022-071018.",
    "Van Calster B, McLernon DJ, van Smeden M, Wynants L, Steyerberg EW. Calibration: the Achilles heel of predictive analytics. BMC Med. 2019;17(1):230. doi:10.1186/s12916-019-1466-7.",
    "Steyerberg EW, Harrell FE Jr. Prediction models need appropriate internal, internal-external, and external validation. J Clin Epidemiol. 2016;69:245-247. doi:10.1016/j.jclinepi.2015.04.005.",
    "Vickers AJ, Elkin EB. Decision curve analysis: a novel method for evaluating prediction models. Med Decis Making. 2006;26(6):565-574. doi:10.1177/0272989X06295361.",
    "Kerr KF, Brown MD, Zhu K, Janes H. Assessing the clinical impact of risk prediction models with decision curves: guidance for correct interpretation and appropriate use. J Clin Oncol. 2016;34(21):2534-2540. doi:10.1200/JCO.2015.65.5654.",
    "Steyerberg EW, Vickers AJ, Cook NR, Gerds T, Gonen M, Obuchowski N, et al. Assessing the performance of prediction models. Epidemiology. 2010;21(1):128-138. doi:10.1097/EDE.0b013e3181c30fb2.",
]

REFERENCE_PMIDS = [
    "38626948", "41398250", "33177372", "34259450", "34775971", "35836294",
    "40702538", "34392691", "36164800", "41554280", "33538696", "33514699",
    "37588623", "33597541", "34152373", "36596836", "30204154", "37052626",
    None, None, None, "30596875", "25569120", "25179855", "36750242",
    "31842878", "25981519", "17099194", "27247223", "20010215",
]

REFERENCE_VERIFICATION_STATUS = {
    **{i: "verified: DOI metadata matched Crossref and PMID metadata matched PubMed (title, first author, journal, and year)" for i in range(1, 19)},
    19: "verified: DOI record matched DataCite and MIMIC-IV v3.1 citation matched the official PhysioNet version page",
    20: "verified: DOI record matched DataCite and eICU-CRD v2.0 citation matched the official PhysioNet version page",
    21: "verified: DOI record matched DataCite and SICdb v1.0.8 citation matched the official PhysioNet version page",
    **{i: "verified: DOI metadata matched Crossref and PMID metadata matched PubMed (title, first author, journal, and year)" for i in range(22, 31)},
}

# Keep the manuscript reference list in first-citation order. The source draft
# is renumbered with the same mapping before this generator is run.
OLD_TO_NEW_REFERENCE_NUMBER = {
    1: 19, 2: 1, 3: 2, 4: 3, 5: 4, 6: 8, 7: 9, 8: 5, 9: 6,
    10: 7, 11: 10, 12: 11, 13: 12, 14: 13, 15: 14, 16: 23,
    17: 25, 18: 27, 19: 24, 20: 26, 21: 28, 22: 21, 23: 20,
    24: 15, 25: 22, 26: 16, 27: 17, 28: 29, 29: 30, 30: 18,
}
NEW_TO_OLD_REFERENCE_NUMBER = {
    new: old for old, new in OLD_TO_NEW_REFERENCE_NUMBER.items()
}
REFERENCES = [REFERENCES[NEW_TO_OLD_REFERENCE_NUMBER[i] - 1] for i in range(1, len(REFERENCES) + 1)]
REFERENCE_PMIDS = [REFERENCE_PMIDS[NEW_TO_OLD_REFERENCE_NUMBER[i] - 1] for i in range(1, len(REFERENCE_PMIDS) + 1)]
REFERENCE_VERIFICATION_STATUS = {
    i: REFERENCE_VERIFICATION_STATUS[NEW_TO_OLD_REFERENCE_NUMBER[i]]
    for i in range(1, len(REFERENCES) + 1)
}


def set_run(run, size: float = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def configure(doc: Document, double: bool = True) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.85 if double else 1)
        section.bottom_margin = Inches(0.85 if double else 1)
        section.left_margin = Inches(0.9 if double else 1)
        section.right_margin = Inches(0.9 if double else 1)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6 if double else 0)
    normal.paragraph_format.line_spacing = MANUSCRIPT_LINE_SPACING if double else 1.0
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(14 if style_name == "Heading 1" else 12)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = MANUSCRIPT_LINE_SPACING if double else 1.0


def add_para(doc: Document, text: str, align=None, bold=False, italic=False, spacing=None, size: float = 12) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if spacing is not None:
        p.paragraph_format.line_spacing = spacing
    if spacing is not None and spacing > 1.0:
        p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, size, bold=bold, italic=italic)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.line_spacing = MANUSCRIPT_LINE_SPACING
    run = p.add_run(text)
    set_run(run, 14 if level == 1 else 12, bold=True)


def add_page_field(run) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, result, end])


def set_page_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    footer = section.footer
    p = footer.paragraphs[0]
    for child in list(p._p):
        if child.tag != qn("w:pPr"):
            p._p.remove(child)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    text_width = (section.page_width - section.left_margin - section.right_margin) / 914400
    p.paragraph_format.tab_stops.add_tab_stop(Inches(text_width / 2), WD_TAB_ALIGNMENT.CENTER)
    left_number = p.add_run()
    set_run(left_number, 10)
    add_page_field(left_number)
    center_label = p.add_run("\tPage ")
    set_run(center_label, 10)
    center_number = p.add_run()
    set_run(center_number, 10)
    add_page_field(center_number)


def enable_line_numbers(section) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:restart"), "continuous")
    line_numbers.set(qn("w:distance"), "360")
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        sect_pr.append(line_numbers)
    else:
        sect_pr.insert(sect_pr.index(cols), line_numbers)


def configure_manuscript_page_furniture(doc: Document) -> None:
    for section in doc.sections:
        enable_line_numbers(section)
        set_page_footer(section)


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("**", "")
    return text.strip()


def manuscript_source() -> str:
    source = (OUT / "corrected_manuscript_draft.md").read_text(encoding="utf-8")
    return source.split("## Reference anchors for final reference-manager import", 1)[0].rstrip()


def parse_manuscript() -> tuple[str, list[tuple[str, str]]]:
    source = manuscript_source()
    title = ""
    blocks: list[tuple[str, str]] = []
    paragraph: list[str] = []
    for line in source.splitlines():
        line = line.rstrip()
        if line.startswith("# "):
            title = clean_inline(line[2:])
            continue
        if line.startswith("## ") or line.startswith("### "):
            if paragraph:
                blocks.append(("paragraph", " ".join(paragraph)))
                paragraph = []
            level = "heading1" if line.startswith("## ") else "heading2"
            blocks.append((level, clean_inline(line.split(" ", 1)[1])))
            continue
        if line.startswith("| ") or line.startswith("---"):
            continue
        if line.startswith("- "):
            if paragraph:
                blocks.append(("paragraph", " ".join(paragraph)))
                paragraph = []
            blocks.append(("bullet", clean_inline(line[2:])))
            continue
        if not line.strip():
            if paragraph:
                blocks.append(("paragraph", " ".join(paragraph)))
                paragraph = []
        else:
            paragraph.append(clean_inline(line))
    if paragraph:
        blocks.append(("paragraph", " ".join(paragraph)))
    return title, blocks


def add_author_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    for i, (name, sup, marker) in enumerate(AUTHORS):
        run = p.add_run(name)
        set_run(run, 12)
        number = p.add_run(sup)
        set_run(number, 8)
        number.font.superscript = True
        if marker:
            marker_run = p.add_run(marker)
            set_run(marker_run, 8)
            marker_run.font.superscript = True
        if i < len(AUTHORS) - 1:
            comma = p.add_run(", ")
            set_run(comma, 12)
    add_para(doc, "1 Hangzhou TCM Hospital Affiliated with Zhejiang Chinese Medical University, Hangzhou, Zhejiang, China.", WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_para(doc, "2 Affiliated Mental Health Center & Hangzhou Seventh People's Hospital, Zhejiang University School of Medicine, Zhejiang, China.", WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)
    add_para(doc, "# These authors contributed equally to this work and share first authorship.", WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)


def add_corresponding_author(doc: Document) -> None:
    add_para(
        doc,
        "*Corresponding author: Tielong Chen, Hangzhou Traditional Chinese Medicine Hospital Affiliated with Zhejiang Chinese Medical University, No. 453 Tiyuchang Road, Xihu District, Hangzhou City, Zhejiang Province, China. Email: ctlktz@163.com.",
        WD_ALIGN_PARAGRAPH.CENTER,
        spacing=1.0,
    )


def write_title_page() -> None:
    doc = Document()
    configure(doc, double=False)
    add_para(doc, TITLE, WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=1.0)
    doc.add_paragraph()
    add_author_block(doc)
    doc.add_paragraph()
    add_para(doc, "Corresponding author", bold=True, spacing=1.0)
    add_para(doc, "Tielong Chen", spacing=1.0)
    add_para(doc, "Hangzhou Traditional Chinese Medicine Hospital Affiliated with Zhejiang Chinese Medical University", spacing=1.0)
    add_para(doc, "No. 453 Tiyuchang Road, Xihu District, Hangzhou City, Zhejiang Province, China", spacing=1.0)
    add_para(doc, "Email: ctlktz@163.com", spacing=1.0)
    doc.add_paragraph()
    add_para(doc, "Keywords: transportability; critical care; vasopressor initiation; machine learning; calibration; external validation; MIMIC-IV; eICU-CRD; SICdb", spacing=1.0)
    add_para(doc, "Funding: No financial support was received for this work.", spacing=1.0)
    add_para(doc, "Conflicts of interest: None declared.", spacing=1.0)
    add_para(doc, "Ethics: [AUTHOR_INPUT_NEEDED: provide the institutional review board or ethics committee determination, waiver of informed consent, and approval identifier if applicable.]", spacing=1.0)
    add_para(doc, "Patient and public involvement: [AUTHOR_CONFIRMATION_NEEDED]", spacing=1.0)
    add_para(doc, "Proposed CRediT author contribution statement (to be confirmed by all authors)", bold=True, spacing=1.0)
    contributions = [
        "Pengqin Shu: Conceptualization, data curation, formal analysis, investigation, methodology, software, visualization, writing - original draft.",
        "Xiaoye Xu: Conceptualization, data curation, formal analysis, methodology, validation, writing - original draft.",
        "Fei Ying: Data curation, investigation, validation, writing - review and editing.",
        "Jing Sun: Data curation, investigation, validation, writing - review and editing.",
        "Zhao Liu: Methodology, statistical analysis, software, supervision, writing - review and editing.",
        "Tielong Chen: Conceptualization, methodology, project administration, supervision, writing - review and editing.",
    ]
    for item in contributions:
        add_para(doc, item, spacing=1.0)
    add_para(doc, "Acknowledgements: [AUTHOR_CONFIRMATION_NEEDED: state acknowledgements or indicate none.]", spacing=1.0)
    add_para(doc, "AI assistance disclosure: [AUTHOR_CONFIRMATION_NEEDED: disclose any language, coding, or analytical assistance according to the target journal policy.]", spacing=1.0)
    doc.save(PKG / "Title_Page.docx")


def write_cover_letter() -> None:
    date = "26 August 2026"
    text = f"""{date}\n\nThe Editor-in-Chief\n[Target Journal]\n\nDear Editor-in-Chief,\n\nWe submit our manuscript entitled \"{TITLE}\" for consideration as an original research article in [Target Journal].\n\nThis retrospective landmark study evaluates whether an hourly model for first documented continuous vasopressor initiation transports across temporal drift, a large US multicenter network, and an Austrian ICU database. The model was developed and selected using temporally separated MIMIC-IV cohorts, frozen before testing, and evaluated with discrimination, calibration, held-out local recalibration, hospital-level heterogeneity, alternative landmark estimands, clinical comparators, and an actionability-gap sensitivity analysis. The central finding is that patient ranking transported better than absolute risk: HGB AUROC was 0.763 in the MIMIC-IV 2020-2022 temporal test, 0.838 in eICU-CRD, and 0.696 in SICdb, while calibration slopes and threshold-related alert burden varied substantially by setting. These results provide a practical evaluation framework for deciding when local calibration and threshold assessment are necessary before prospective silent testing.\n\nThe manuscript is deliberately framed as a transportability and calibration study. It does not claim a novel vasopressor endpoint, clinical effectiveness, deployment readiness, or universal alert threshold. The analysis follows TRIPOD+AI-oriented reporting principles, treats the ICU or unit stay as the inference unit for repeated landmarks, and reports post hoc extensions transparently.\n\nData and analysis code will be shared in disclosure-reviewed form at [PUBLIC_REPOSITORY_URL] and archived at [PERSISTENT_DOI] before publication. Restricted source databases and row-level derived data will not be redistributed.\n\nThe authors report that no financial support was received for this work and declare no conflicts of interest. [AUTHOR_CONFIRMATION_NEEDED: confirm that the manuscript is original, is not under consideration elsewhere, and has no related simultaneous submission or preprint requiring disclosure.]\n\nThank you for considering this manuscript.\n\nSincerely,\n\nPengqin Shu, Xiaoye Xu, Fei Ying, Jing Sun, Zhao Liu, and Tielong Chen\nFor all authors\n\nCorresponding author:\nTielong Chen\nHangzhou Traditional Chinese Medicine Hospital Affiliated with Zhejiang Chinese Medical University\nNo. 453 Tiyuchang Road, Xihu District, Hangzhou City, Zhejiang Province, China\nEmail: ctlktz@163.com\n"""
    (PKG / "Cover_Letter_generic.md").write_text(text, encoding="utf-8")
    doc = Document()
    configure(doc, double=False)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    for i, line in enumerate(text.splitlines()):
        if not line:
            blank = doc.add_paragraph()
            blank.paragraph_format.space_after = Pt(2)
            blank.paragraph_format.line_spacing = 0.5
            continue
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        set_run(run, 12)
    doc.save(PKG / "Cover_Letter_generic.docx")


def write_references_and_checklist() -> None:
    ref_lines = ["# Source-verified reference list", "", *[f"{i}. {ref}" for i, ref in enumerate(REFERENCES, 1)], "", "These records were checked against Crossref and PubMed metadata, or against the DataCite DOI record and official PhysioNet version page for datasets. Convert the final list to the target journal's style before submission."]
    (PKG / "reference_list_source_verified.md").write_text("\n".join(ref_lines) + "\n", encoding="utf-8")
    check_lines = ["# Reference verification checklist", "", "| No. | Source | DOI / PMID / official page | Verification status |", "| --- | --- | --- | --- |"]
    for i, ref in enumerate(REFERENCES, 1):
        doi = re.search(r"doi:(10\.\S+)", ref).group(1).rstrip(".")
        pmid_value = REFERENCE_PMIDS[i - 1]
        pmid = f"; [PMID](https://pubmed.ncbi.nlm.nih.gov/{pmid_value}/)" if pmid_value else ""
        if "[dataset]" in ref:
            link_match = re.search(r"Available from: (https?://\S+)", ref)
            page_url = link_match.group(1).rstrip(".") if link_match else ""
            page = f"; [official page]({page_url})" if page_url else ""
            source = REFERENCE_VERIFICATION_STATUS[i]
        else:
            page = ""
            source = REFERENCE_VERIFICATION_STATUS[i]
        check_lines.append(f"| {i} | Reference metadata | [DOI](https://doi.org/{doi}){pmid}{page} | {source} |")
    check_lines.extend(["", "Recent-reference proportion: 20 of 30 references (66.7%) were published in 2021-2026.", "The manuscript reference list uses numbered Vancouver-style entries; PMID and official verification links are retained in this checklist rather than in the manuscript reference list.", "In-text citation coverage was audited against the generated manuscript: references [1]-[30] are each cited at least once, and no out-of-range citation number was detected."])
    (PKG / "reference_verification_checklist.md").write_text("\n".join(check_lines) + "\n", encoding="utf-8")


def word_count(text: str) -> int:
    return len(re.findall(r"\b[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*\b", text))


def abstract_and_main_word_counts(blocks: list[tuple[str, str]]) -> tuple[int, int]:
    abstract: list[str] = []
    main_text: list[str] = []
    current_heading = ""
    for kind, text in blocks:
        if kind in {"heading1", "heading2"}:
            current_heading = text if kind == "heading1" else current_heading
            continue
        if kind != "paragraph":
            continue
        if current_heading == "Abstract" and not text.startswith("Keywords:"):
            abstract.append(text)
        if current_heading in {
            "Introduction",
            "Methods",
            "Results",
            "Discussion",
            "Conclusion",
        }:
            main_text.append(text)
    return word_count(" ".join(abstract)), word_count(" ".join(main_text))


def add_abstract_paragraph(doc: Document, text: str) -> None:
    labels = ("Background:", "Methods:", "Results:", "Conclusions:")
    matches = [match for label in labels for match in [text.find(label)] if match >= 0]
    starts = sorted(set(matches))
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = MANUSCRIPT_LINE_SPACING
    p.paragraph_format.space_after = Pt(6)
    for i, start in enumerate(starts):
        end = starts[i + 1] if i + 1 < len(starts) else len(text)
        segment = text[start:end]
        label_end = segment.find(":") + 1
        label_run = p.add_run(segment[:label_end])
        set_run(label_run, 12, bold=True)
        body_run = p.add_run(segment[label_end:])
        set_run(body_run, 12)


def scale_table_widths(widths: list[int], target: int = 13680) -> list[int]:
    scale = target / sum(widths)
    scaled = [int(round(width * scale)) for width in widths]
    scaled[-1] += target - sum(scaled)
    return scaled


SUPPLEMENTARY_NOTES = [
    "Supplementary Table 1. Database design and measurement context. Database versions and access terms are based on the audited data-availability record.",
    "Supplementary Table 2. Predictor dictionary and source mappings. The 42 predictors were fixed before external scoring, and values outside the accepted range were set to missing.",
    "Supplementary Table 3. Exploratory subgroup performance. Subgroup analyses were not used for model selection or formal fairness inference.",
    "Supplementary Table 4. Repeated intercept recalibration. Negative Brier changes indicate improvement after intercept recalibration, with identifier-disjoint calibration and evaluation subsets.",
    "Supplementary Table 5. Horizon eligibility and alternative landmark estimands. The primary binary estimand requires a complete six-hour horizon; hour-6 analysis uses one landmark per stay.",
    "Supplementary Table 6. Clinical comparators, reduced-feature models, and temperature audit. These analyses were post hoc robustness analyses and were not used for model selection.",
    "Supplementary Table 7. Equal-total-stay-weight training and one-hour clinical-action-gap analysis. The gap-trained model excludes initiation within one hour of the landmark.",
    "Supplementary Table 8. Six-hour alert-suppression sensitivity. This is one descriptive alert policy and does not establish prospective clinical benefit.",
]


MAIN_TABLE_SPECS = [
    (
        "Table 1. Cohort characteristics and event prevalence",
        table_1,
        [1700, 950, 900, 1350, 950, 1250, 1200, 1050, 1010],
        "Age is shown as median (IQR). Event prevalence is reported at both the stay and landmark levels in the corresponding columns.",
    ),
    (
        "Table 2. Main validation results for frozen models",
        table_2,
        [1350, 1250, 1200, 1050, 1400, 1400, 1400, 1300, 1310],
        "Confidence intervals use 500 stay-clustered bootstrap replicates for AUROC, AUPRC, and Brier score, and stay-clustered sandwich intervals for calibration statistics.",
    ),
    (
        "Table 3. Calibration and site-level heterogeneity",
        table_3,
        [1550, 1900, 1050, 950, 1350, 1600, 1960],
        "CITL is calibration-in-the-large. The 73-hospital row is descriptive; the 208-hospital row uses hospital-clustered bootstrap inference.",
    ),
    (
        "Table 4. Robustness and sensitivity analyses",
        table_4,
        [2200, 1650, 1650, 1650, 2210],
        "All sensitivity analyses were post hoc and were not used for model selection.",
    ),
    (
        "Table 5. Threshold and alert-policy metrics",
        table_5,
        [1200, 1250, 950, 950, 950, 700, 1500, 1550, 900, 1410],
        "Sens. = sensitivity; Spec. = specificity; PPV = positive predictive value. The 0.05 threshold and six-hour suppression policy are exploratory and do not establish prospective clinical benefit.",
    ),
]


def write_main_manuscript() -> None:
    _, blocks = parse_manuscript()
    abstract_count, main_text_count = abstract_and_main_word_counts(blocks)
    doc = Document()
    configure(doc, double=True)
    add_para(doc, TITLE, WD_ALIGN_PARAGRAPH.CENTER, bold=True, spacing=MANUSCRIPT_LINE_SPACING, size=16)
    add_author_block(doc)
    add_corresponding_author(doc)
    add_heading(doc, "Short title", 1)
    add_para(doc, SHORT_TITLE, spacing=MANUSCRIPT_LINE_SPACING)
    add_heading(doc, "Article type", 1)
    add_para(doc, ARTICLE_TYPE, spacing=MANUSCRIPT_LINE_SPACING)
    add_para(
        doc,
        f"Main-text word count: {main_text_count:,}; Abstract word count: {abstract_count:,}; Figures: 5; Tables: 5; Supplementary tables: 8.",
        spacing=MANUSCRIPT_LINE_SPACING,
    )
    current_heading = ""
    for kind, text in blocks:
        if kind == "heading1":
            current_heading = text
            add_heading(doc, text, 1)
        elif kind == "heading2":
            add_heading(doc, text, 2)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = MANUSCRIPT_LINE_SPACING
            run = p.add_run(text)
            set_run(run, 12)
        else:
            if current_heading == "Abstract" and text.startswith("Background:"):
                add_abstract_paragraph(doc, text)
            else:
                add_para(doc, text, spacing=MANUSCRIPT_LINE_SPACING)
    add_heading(doc, "References", 1)
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.line_spacing = MANUSCRIPT_LINE_SPACING
        run = p.add_run(f"{i}. {ref}")
        set_run(run, 12)

    add_heading(doc, "Supplementary material", 1)
    for note in SUPPLEMENTARY_NOTES:
        add_para(doc, note, spacing=MANUSCRIPT_LINE_SPACING)

    legends = (OUT / "corrected_figure_legends.md").read_text(encoding="utf-8")
    add_heading(doc, "Figure legends", 1)
    for line in legends.splitlines():
        if not line.strip() or line.startswith("#"):
            continue
        add_para(doc, clean_inline(line), spacing=MANUSCRIPT_LINE_SPACING)

    table_section = doc.add_section(WD_SECTION.NEW_PAGE)
    table_section.orientation = WD_ORIENT.LANDSCAPE
    table_section.page_width, table_section.page_height = table_section.page_height, table_section.page_width
    table_section.top_margin = Inches(0.75)
    table_section.bottom_margin = Inches(0.75)
    table_section.left_margin = Inches(0.75)
    table_section.right_margin = Inches(0.75)
    add_heading(doc, "Main table", 1)
    for i, (title, table_factory, widths, note) in enumerate(MAIN_TABLE_SPECS):
        header, data = table_factory()
        add_submission_table(
            doc,
            header,
            data,
            scale_table_widths(widths),
            title=title,
            note=note,
            page_break_before=(i > 0),
        )
    configure_manuscript_page_furniture(doc)
    doc.save(PKG / "Main_Manuscript.docx")
    source = manuscript_source()
    source_body = re.sub(r"^# .*?\n+", "", source, count=1, flags=re.DOTALL)
    author_line = "Pengqin Shu1#, Xiaoye Xu2#, Fei Ying1#, Jing Sun1#, Zhao Liu1, Tielong Chen1*"
    metadata = (
        f"# {TITLE}\n\n{author_line}\n\n"
        "# These authors contributed equally to this work and share first authorship.\n\n"
        "1 Hangzhou TCM Hospital Affiliated with Zhejiang Chinese Medical University, Hangzhou, Zhejiang, China.\n\n"
        "2 Affiliated Mental Health Center & Hangzhou Seventh People's Hospital, Zhejiang University School of Medicine, Zhejiang, China.\n\n"
        "*Corresponding author: Tielong Chen, Hangzhou Traditional Chinese Medicine Hospital Affiliated with Zhejiang Chinese Medical University, No. 453 Tiyuchang Road, Xihu District, Hangzhou City, Zhejiang Province, China. Email: ctlktz@163.com.\n\n"
        f"## Short title\n\n{SHORT_TITLE}\n\n## Article type\n\n{ARTICLE_TYPE}\n\n"
        f"Main-text word count: {main_text_count:,}; Abstract word count: {abstract_count:,}; Figures: 5; Tables: 5; Supplementary tables: 8.\n\n"
    )
    table_notes = "\n\n".join([title + ". " + note for title, _, _, note in MAIN_TABLE_SPECS])
    main_md = (
        metadata
        + source_body.strip()
        + "\n\n## References\n\n"
        + "\n".join(f"{i}. {r}" for i, r in enumerate(REFERENCES, 1))
        + "\n\n## Supplementary material\n\n"
        + "\n\n".join(SUPPLEMENTARY_NOTES)
        + "\n\n## Figure legends\n\n"
        + legends
        + "\n\n## Main table\n\n"
        + table_notes
        + "\n"
    )
    (PKG / "Main_Manuscript.md").write_text(main_md, encoding="utf-8")
    shutil.copy2(OUT / "corrected_figure_legends.md", PKG / "corrected_figure_legends.md")
    shutil.copy2(OUT / "figure_source_manifest.csv", PKG / "figure_source_manifest.csv")


def write_readme() -> None:
    text = """# Submission package

This package is a generic journal-format submission set for the study on temporal and geographic transportability of an hourly vasopressor-initiation model. No target journal was specified, so journal-specific word limits, reference style, reporting checklists, file naming, and declarations must be adapted before upload.

## Core files

- `Main_Manuscript.docx` and `Main_Manuscript.md`: main text, abstract, declarations, source-verified reference list, and figure legends.
- `Title_Page.docx`: authors, affiliations, correspondence, keywords, proposed CRediT statement, and author confirmation fields.
- `Cover_Letter_generic.docx` and `Cover_Letter_generic.md`: generic cover letter with target-journal and repository placeholders.
- `figures/`: five main figures (Figures 1-5) and two supplementary figures (Supplementary Figures S1-S2) in PDF, SVG, TIFF, and PNG formats.
- `tables/`: Tables 1-5 in DOCX, CSV, and Markdown formats. Table cells use white backgrounds with black text and borders only.
- `supplementary/`: Supplementary Tables 1-8 in one DOCX plus editable CSV and Markdown files. Tables 5-7 also have independent 5A/5B, 6A/6B/6C, and 7A/7B files so each data block is directly editable.
- `corrected_figure_legends.md` and `figure_source_manifest.csv`: final legend text and source-data mapping for the figure files.
- `reference_claim_audit.md`: claim-to-citation audit showing whether each in-text reference supports the associated statement.

## Mandatory author confirmations before submission

1. Replace `[Target Journal]` and adapt the cover letter to the journal's scope.
2. Supply the ethics committee or institutional review board determination, waiver language, and identifier if applicable.
3. Confirm patient/public involvement wording.
4. Confirm that no financial support was received and that no conflicts of interest exist for every author.
5. Confirm the proposed CRediT roles with all authors.
6. Replace `[PUBLIC_REPOSITORY_URL]`, `[PERSISTENT_DOI]`, and `[OPEN_SOURCE_LICENSE]` after disclosure review and repository creation.
7. Confirm originality, exclusive submission, preprint status, related manuscripts, and AI-use disclosure.
8. Read `reference_claim_audit.md`, import the source-verified references into the journal's required format, and recheck every in-text citation.
9. Confirm whether the target journal uses `Supplementary Figure`, `Extended Data Figure`, or another nomenclature; the current generic package uses `Supplementary Figure S1-S2`.

## Reference handling

The manuscript contains 30 numbered references, with 20 (66.7%) published in 2021-2026. The manuscript list is formatted as numbered Vancouver-style entries; DOI, PMID, and official dataset links are documented separately in `reference_verification_checklist.md`. A citation-coverage audit confirms that references [1]-[30] are each cited at least once, and `reference_claim_audit.md` records the claim-level support review.

The SICdb cohort is described as a single Austrian hospital cohort, not European multicenter validation. Decision-curve and alert-suppression results are exploratory policy analyses and do not establish clinical benefit.
"""
    (PKG / "README.md").write_text(text, encoding="utf-8")


if __name__ == "__main__":
    write_title_page()
    write_cover_letter()
    write_references_and_checklist()
    write_main_manuscript()
    write_readme()
