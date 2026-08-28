"""Generate editable submission tables and Word table packages from audited CSV outputs."""

from __future__ import annotations

import csv
import math
import statistics
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
PKG = ROOT / "submission_package_2026-08-26"
TABLES = PKG / "tables"
SUPP = PKG / "supplementary"
TABLES.mkdir(parents=True, exist_ok=True)
SUPP.mkdir(parents=True, exist_ok=True)

FONT = "Times New Roman"
TOTAL_DXA = 9360


def rows(name: str) -> list[dict[str, str]]:
    with (OUT / name).open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def write_csv(path: Path, header: list[str], data: list[list[str]]) -> None:
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(header)
        writer.writerows(data)


def write_markdown(path: Path, title: str, header: list[str], data: list[list[str]], note: str = "") -> None:
    lines = [f"# {title}", "", "| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in data)
    if note:
        lines.extend(["", "Note: " + note])
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def f(value: str | float, digits: int = 3) -> str:
    return f"{float(value):.{digits}f}"


def pct(value: str | float, digits: int = 1) -> str:
    return f"{100 * float(value):.{digits}f}%"


def integer(value: str | float) -> str:
    return f"{int(round(float(value))):,}"


def ci(row: dict[str, str], metric: str, digits: int = 3) -> str:
    return f"{f(row[metric], digits)} ({f(row[metric + '_ci_low'], digits)} to {f(row[metric + '_ci_high'], digits)})"


def ci_clustered(row: dict[str, str], metric: str, digits: int = 3) -> str:
    return f"{f(row[metric], digits)} ({f(row[metric + '_ci_low'], digits)} to {f(row[metric + '_ci_high'], digits)})"


def label_dataset(name: str) -> str:
    return {
        "mimic_development_2008_2016": "MIMIC-IV development, 2008-2016",
        "mimic_model_selection_2017_2019": "MIMIC-IV model selection, 2017-2019",
        "mimic_temporal_test_2020_2022": "MIMIC-IV temporal test, 2020-2022",
        "mimic_temporal_test": "MIMIC-IV temporal test, 2020-2022",
        "eicu_external": "eICU-CRD external validation",
        "sicdb_external": "SICdb external validation",
        "mimiciv": "MIMIC-IV",
        "eicu": "eICU-CRD",
        "sicdb": "SICdb",
    }.get(name, name)


def model_label(name: str) -> str:
    return {
        "logistic_regression": "Regularized logistic regression",
        "hist_gradient_boosting": "Histogram-based gradient boosting",
        "full_hgb": "Full HGB",
        "vital_only_hgb": "Vital-sign-only HGB",
        "full_minus_temperature_hgb": "Full HGB without temperature",
        "primary_landmark_weighted_hgb": "Primary HGB",
        "stay_weighted_training_hgb": "Equal-total-stay-weight HGB",
        "primary_six_hour_hgb_on_gap_cohort": "Primary HGB on gap-eligible rows",
        "one_hour_gap_trained_hgb": "One-hour-gap trained HGB",
    }.get(name, name)


def set_run_font(run, size: float = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 90) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def set_borders(table, color: str = "000000", size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_header_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)


def configure_document(doc: Document, landscape: bool = False) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(12)


def add_title(doc: Document, text: str, page_break_before: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = page_break_before
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, 12, bold=True)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, 10)


def add_table(
    doc: Document,
    header: list[str],
    data: list[list[str]],
    widths: list[int],
    title: str | None = None,
    note: str = "",
    page_break_before: bool = False,
) -> None:
    if title:
        add_title(doc, title, page_break_before=page_break_before)
    table = doc.add_table(rows=1, cols=len(header))
    table.allow_autofit = False
    for i, value in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(value)
        set_run_font(run, 9.0, bold=True)
    set_header_repeat(table.rows[0])
    set_row_cant_split(table.rows[0])
    for ridx, row in enumerate(data):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(value))
            set_run_font(run, 9.0)
        set_row_cant_split(table.rows[-1])
    set_table_geometry(table, widths)
    set_borders(table)
    if note:
        add_note(doc, note)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def save_table_doc(
    path: Path,
    title: str,
    sections: list[tuple[list[str], list[list[str]], list[int], str]],
    section_titles: list[str] | None = None,
    page_break_between: bool = False,
) -> None:
    doc = Document()
    landscape = any(len(header) >= 8 for header, _, _, _ in sections)
    configure_document(doc, landscape=landscape)
    for i, (header, data, widths, note) in enumerate(sections):
        # Use the full landscape text width so supplementary tables do not
        # leave a large unused margin or force avoidable mid-word wrapping.
        if landscape:
            target_width = 13680
            scale = target_width / sum(widths)
            scaled = [int(round(width * scale)) for width in widths]
            scaled[-1] += target_width - sum(scaled)
            widths = scaled
        current_title = title if i == 0 else None
        if section_titles is not None:
            current_title = section_titles[i]
        add_table(
            doc,
            header,
            data,
            widths,
            current_title,
            note,
            page_break_before=(i > 0 and page_break_between),
        )
    doc.save(path)


def q1q3(r: dict[str, str]) -> str:
    return f"{f(r['age_median'], 0)} ({f(r['age_q1'], 0)}-{f(r['age_q3'], 0)})"


def table_1() -> tuple[list[str], list[list[str]]]:
    header = ["Cohort", "Stays", "Patients", "Age, median (IQR)", "Male stays (%)", "Event-positive stays (n)", "Event-positive stays (%)", "Landmarks (n)", "Positive landmarks (%)"]
    data = []
    for r in rows("corrected_table1_cohort_summary.csv"):
        data.append([
            label_dataset(r["dataset"]), integer(r["n_stays"]), integer(r["n_patients"]), q1q3(r),
            pct(float(r["male_percent"]) / 100), integer(r["event_positive_stays"]), pct(float(r["event_positive_stays_percent"]) / 100),
            integer(r["n_landmarks"]), pct(float(r["positive_landmarks_percent"]) / 100),
        ])
    return header, data


def table_2() -> tuple[list[str], list[list[str]]]:
    header = ["Cohort", "Model", "Landmarks / stays", "Positive landmarks", "AUROC (95% CI)", "AUPRC (95% CI)", "Brier (95% CI)", "CITL (95% CI)", "Calibration slope (95% CI)"]
    data = []
    for r in rows("corrected_clustered_confidence_intervals.csv"):
        data.append([
            label_dataset(r["dataset"]), model_label(r["model"]), f"{integer(r['n_samples'])} / {integer(r['n_clusters'])}",
            pct(r["n_positive_records"] if "n_positive_records" in r else "0"), ci_clustered(r, "auroc"), ci_clustered(r, "auprc"),
            ci_clustered(r, "brier", 5), ci_clustered(r, "calibration_in_the_large", 3), ci_clustered(r, "calibration_slope", 3),
        ])
    # The source table uses n_positive_records only in the validation report, so restore the observed prevalence from the frozen metrics.
    frozen = {(r["dataset"], r["model"]): r for r in rows("corrected_frozen_validation_metrics.csv")}
    for i, r in enumerate(rows("corrected_clustered_confidence_intervals.csv")):
        data[i][3] = pct(frozen[(r["dataset"], r["model"])]["record_event_prevalence"])
    return header, data


def table_3() -> tuple[list[str], list[list[str]]]:
    header = ["Analysis level", "Cohort / site", "Landmarks", "Clusters", "Brier (95% CI)", "CITL (95% CI)", "Calibration slope (95% CI)"]
    data = []
    pooled = rows("corrected_clustered_confidence_intervals.csv")
    for dataset in ["mimic_temporal_test", "eicu_external", "sicdb_external"]:
        r = next(row for row in pooled if row["dataset"] == dataset and row["model"] == "hist_gradient_boosting")
        data.append([
            "Pooled stay-clustered", label_dataset(dataset), integer(r["n_samples"]), integer(r["n_clusters"]),
            ci_clustered(r, "brier", 5), ci_clustered(r, "calibration_in_the_large"), ci_clustered(r, "calibration_slope"),
        ])
    hospital = next(row for row in rows("corrected_hierarchical_cluster_metrics.csv") if row["cluster_level"] == "hospital_id")
    data.append([
        "Hospital bootstrap", "eICU-CRD (208 hospitals)", integer(hospital["n_landmarks"]), integer(hospital["n_clusters"]),
        ci_clustered(hospital, "brier", 5), ci_clustered(hospital, "calibration_in_the_large"), ci_clustered(hospital, "calibration_slope"),
    ])
    summary = {row["metric"]: row for row in rows("corrected_eicu_hospital_calibration_summary.csv")}
    citl = summary["calibration_in_the_large"]
    slope = summary["calibration_slope"]
    data.append([
        "Descriptive hospital range", "eICU-CRD (73 hospitals)", "Not applicable", "73",
        "Not estimated", f"Median {float(citl['median']):.3f} (range {float(citl['minimum']):.3f} to {float(citl['maximum']):.3f})",
        f"Median {float(slope['median']):.3f} (range {float(slope['minimum']):.3f} to {float(slope['maximum']):.3f})",
    ])
    return header, data


def table_4() -> tuple[list[str], list[list[str]]]:
    header = ["Sensitivity analysis", "MIMIC-IV 2020-2022 AUROC (95% CI)", "eICU-CRD AUROC (95% CI)", "SICdb AUROC (95% CI)", "Interpretation"]
    data = []
    primary = {r["dataset"]: r for r in rows("corrected_clustered_confidence_intervals.csv") if r["model"] == "hist_gradient_boosting"}
    fixed = {r["dataset"]: r for r in rows("corrected_fixed_hour_metrics.csv") if r["index_hour"] == "6"}
    strict = {r["dataset"].replace("_strict_future", ""): r for r in rows("corrected_strict_future_metrics.csv")}
    balanced = {r["dataset"]: r for r in rows("corrected_stay_balanced_metrics.csv")}
    reduced = rows("corrected_reduced_model_metrics.csv")
    vital = {r["dataset"]: r for r in reduced if r["model"] == "vital_only_hgb"}
    no_temp = {r["dataset"]: r for r in reduced if r["model"] == "full_minus_temperature_hgb"}
    gap = {r["dataset"]: r for r in rows("corrected_one_hour_gap_metrics.csv") if r["model"] == "one_hour_gap_trained_hgb"}
    datasets = ["mimic_temporal_test", "eicu_external", "sicdb_external"]

    def values(source: dict[str, dict[str, str]]) -> list[str]:
        return [ci(source[d], "auroc") if "auroc_ci_low" in source[d] else f(source[d]["auroc"], 3) + " (descriptive)" for d in datasets]

    data.append(["Primary conditional 6-hour estimand", *values(primary), "Reference analysis"])
    data.append(["Single landmark at ICU hour 6", *values(fixed), "Removes repeated-landmark weighting"])
    data.append(["Strict-future window", *values(strict), "Excludes treatment recorded exactly at the landmark"])
    data.append(["Equal total evaluation weight per stay", *values(balanced), "Tests sensitivity to repeated landmarks"])
    data.append(["Vital-sign-only HGB", *values(vital), "Tests dependence on laboratory measurements"])
    data.append(["Full HGB without temperature", *values(no_temp), "Tests sensitivity to SICdb temperature artifacts"])
    data.append(["One-hour clinical-action-gap HGB", *values(gap), "Tests a more actionable but post hoc target"])
    return header, data


def table_5() -> tuple[list[str], list[list[str]]]:
    header = ["Cohort", "Calibration", "Cutoff", "Sens.", "Spec.", "PPV", "False alerts / 100 patient-hours", "6-h suppression event-stay sensitivity", "Episode PPV", "False episodes / 100 patient-days"]
    threshold_rows = [r for r in rows("corrected_threshold_analysis.csv") if r["threshold"] == "0.05"]
    suppression = {
        (r["dataset"], r["calibration"]): r
        for r in rows("corrected_alert_suppression_metrics.csv")
        if r["threshold"] == "0.05" and r["suppression_hours"] == "6"
    }
    data = []
    for r in threshold_rows:
        s = suppression[(r["dataset"], r["calibration"])]
        data.append([
            label_dataset(r["dataset"]), "Uncalibrated" if r["calibration"] == "uncalibrated" else "Intercept recalibrated", "0.05",
            f(r["sensitivity"]), f(r["specificity"]), f(r["ppv"]), f(r["false_alerts_per_100_patient_hours"], 2),
            f(s["event_stay_sensitivity"]), f(s["episode_ppv"]), f(s["false_alert_episodes_per_100_patient_days"], 2),
        ])
    return header, data


def supplementary_1() -> tuple[list[str], list[list[str]]]:
    header = ["Database", "Release", "Setting / period", "Units or hospitals", "Role in study", "Primary linkage", "Outcome record", "Access / sharing"]
    data = [
        ["MIMIC-IV", "v3.1", "Single US academic medical center; anchor-year groups used for temporal partitions", "1 hospital; adult ICU stays", "Development, model selection, temporal test", "subject_id and stay_id", "mimiciv_derived.vasoactive_agent positive-rate records", "Credentialed PhysioNet access; raw and row-level derivatives not redistributable"],
        ["eICU-CRD", "v2.0", "US admissions, 2014-2015", "208 hospitals; adult unit stays", "External multicenter validation", "patienthealthsystemstayid; hospital identifiers", "Positive-rate infusiondrug records mapped to five target drugs", "Credentialed PhysioNet access; person separation across admissions not guaranteed"],
        ["SICdb", "v1.0.8", "Single Austrian tertiary hospital; admissions, 2013-2021", "CWIN and INBD primary; all four units sensitivity", "External geographic validation", "PatientID; unit", "Continuous-use medication records mapped by prespecified DrugIDs", "Credentialed PhysioNet access plus contributor review; raw and row-level derivatives not redistributable"],
    ]
    return header, data


def supplementary_2() -> tuple[list[str], list[list[str]]]:
    header = ["No.", "Feature", "Definition", "Dom.", "Window", "Unit", "Range", "MIMIC", "eICU", "SICdb", "HGB missing", "Logistic missing"]
    data = []
    for i, r in enumerate(rows("corrected_predictor_dictionary.csv"), 1):
        data.append([str(i), r["feature"], r["summary_rule"], r["domain"], r["lookback_window"], r["unit"], r["physiologic_range"], r["source_mimic"], r["source_eicu"], r["source_sicdb"], r["hgb_missing_handling"], r["logistic_missing_handling"]])
    return header, data


def supplementary_3() -> tuple[list[str], list[list[str]]]:
    header = ["Database", "Subgroup variable", "Subgroup", "Landmarks", "Stays", "Positive stays", "AUROC (95% CI)", "AUPRC (95% CI)", "Brier (95% CI)", "CITL", "Slope"]
    data = []
    for r in rows("corrected_subgroup_metrics.csv"):
        data.append([label_dataset(r["dataset"]), r["subgroup_variable"], r["subgroup"], integer(r["n_landmarks"]), integer(r["n_stays"]), integer(r["n_event_positive_stays"]), ci(r, "auroc"), ci(r, "auprc"), ci(r, "brier", 5), f(r["calibration_in_the_large"], 3), f(r["calibration_slope"], 3)])
    return header, data


def iqr(values: list[float]) -> tuple[float, float, float]:
    values = sorted(values)
    return statistics.median(values), values[max(0, math.ceil(0.25 * len(values)) - 1)], values[max(0, math.ceil(0.75 * len(values)) - 1)]


def supplementary_4() -> tuple[list[str], list[list[str]]]:
    header = ["Database", "Splits", "Calibration groups", "Evaluation groups", "Intercept shift, median (IQR)", "Brier change, median (IQR)", "Splits improved"]
    data = []
    for dataset in ["mimic_temporal_test", "eicu_external", "sicdb_external"]:
        subset = [r for r in rows("corrected_repeated_recalibration.csv") if r["dataset"] == dataset]
        shifts = [float(r["intercept_shift"]) for r in subset]
        changes = [float(r["brier_change_after_minus_before"]) for r in subset]
        sm, sl, su = iqr(shifts)
        bm, bl, bu = iqr(changes)
        improved = sum(v < 0 for v in changes)
        data.append([label_dataset(dataset), str(len(subset)), f"{integer(subset[0]['n_calibration_groups'])}", f"{integer(subset[0]['n_evaluation_groups'])}", f"{sm:.3f} ({sl:.3f}-{su:.3f})", f"{bm:.6f} ({bl:.6f}-{bu:.6f})", f"{improved}/{len(subset)} ({100 * improved / len(subset):.0f}%)"])
    return header, data


def supplementary_5() -> tuple[tuple[list[str], list[list[str]]], tuple[list[str], list[list[str]]]]:
    horizon_header = ["Database", "At-risk landmarks", "Complete 6-h landmarks", "Incomplete 6-h landmarks", "Incomplete", "Observed event among incomplete", "Censored without observed event"]
    horizon_data = []
    for r in rows("corrected_horizon_eligibility_audit.csv"):
        incomplete_percent = float(r["incomplete_landmark_percent"])
        horizon_data.append([label_dataset(r["dataset"]), integer(r["at_risk_landmarks_present_at_index"]), integer(r["complete_six_hour_landmarks"]), integer(r["incomplete_six_hour_landmarks"]), f(incomplete_percent, 1) + "%", integer(r["incomplete_with_observed_event"]), integer(r["incomplete_censored_without_observed_event"])])
    est_header = ["Database", "Estimand", "Landmarks", "Stays", "AUROC (95% CI)", "AUPRC (95% CI)", "Brier (95% CI)"]
    est_data = []
    primary = [r for r in rows("corrected_clustered_confidence_intervals.csv") if r["model"] == "hist_gradient_boosting"]
    for r in primary:
        est_data.append([label_dataset(r["dataset"]), "Primary conditional 6-h estimand", integer(r["n_samples"]), integer(r["n_clusters"]), ci_clustered(r, "auroc"), ci_clustered(r, "auprc"), ci_clustered(r, "brier", 5)])
    for r in rows("corrected_stay_balanced_metrics.csv"):
        est_data.append([label_dataset(r["dataset"]), "Equal total evaluation weight per stay", integer(r["n_landmarks"]), integer(r["n_stays"]), ci(r, "auroc"), ci(r, "auprc"), ci(r, "brier", 5)])
    for r in rows("corrected_fixed_hour_metrics.csv"):
        if r["index_hour"] == "6":
            est_data.append([label_dataset(r["dataset"]), "Single landmark at ICU hour 6", integer(r["n_landmarks"]), "Not applicable", ci(r, "auroc"), ci(r, "auprc"), ci(r, "brier", 5)])
    return (horizon_header, horizon_data), (est_header, est_data)


def supplementary_6() -> tuple[
    tuple[list[str], list[list[str]]],
    tuple[list[str], list[list[str]]],
    tuple[list[str], list[list[str]]],
]:
    comp_header = ["Database", "Comparator / model", "Landmarks", "AUROC (95% CI)", "AUPRC (95% CI)"]
    comp_data = []
    baseline = rows("corrected_clinical_baseline_metrics.csv")
    for r in baseline:
        if r["score"] in {"map_alone", "shock_index", "modified_shock_index", "full_hgb_all_rows"}:
            comp_data.append([label_dataset(r["dataset"]), r["score"].replace("_", " "), integer(r["n_landmarks"]), ci(r, "auroc"), ci(r, "auprc")])
    reduced_header = ["Database", "Reduced-feature model", "Landmarks", "AUROC (95% CI)", "AUPRC (95% CI)", "Brier (95% CI)", "Calibration slope (95% CI)"]
    reduced_data = []
    for r in rows("corrected_reduced_model_metrics.csv"):
        reduced_data.append([label_dataset(r["dataset"]), model_label(r["model"]), integer(r["n_landmarks"]), ci(r, "auroc"), ci(r, "auprc"), ci(r, "brier", 5), ci(r, "calibration_slope")])
    temperature_header = ["Database", "Landmarks", "Temperature available", "Median (deg C)", "1st percentile (deg C)", "99th percentile (deg C)", "Below 32 deg C", "Below 34 deg C", "Above 40 deg C"]
    temperature_data = []
    for r in rows("corrected_temperature_audit.csv"):
        temperature_data.append([
            label_dataset(r["dataset"]),
            integer(r["n_landmarks"]),
            integer(r["n_temperature_available"]),
            f(r["temperature_median"], 2),
            f(r["temperature_q01"], 2),
            f(r["temperature_q99"], 2),
            pct(r["fraction_below_32"], 2),
            pct(r["fraction_below_34"], 2),
            pct(r["fraction_above_40"], 2),
        ])
    return (comp_header, comp_data), (reduced_header, reduced_data), (temperature_header, temperature_data)


def supplementary_7() -> tuple[tuple[list[str], list[list[str]]], tuple[list[str], list[list[str]]]]:
    weight_header = ["Database", "Training model", "Landmarks", "Stays", "AUROC (95% CI)", "Brier (95% CI)", "Calibration slope (95% CI)"]
    weight_data = []
    for r in rows("corrected_stay_weighted_model_metrics.csv"):
        weight_data.append([label_dataset(r["dataset"]), model_label(r["model"]), integer(r["n_landmarks"]), integer(r["n_stays"]), ci(r, "auroc"), ci(r, "brier", 5), ci(r, "calibration_slope")])
    gap_header = ["Database", "Target definition", "Landmarks", "Stays", "Excluded imminent positives", "AUROC (95% CI)", "AUPRC (95% CI)", "Brier (95% CI)"]
    gap_data = []
    for r in rows("corrected_one_hour_gap_metrics.csv"):
        gap_data.append([label_dataset(r["dataset"]), model_label(r["model"]), integer(r["n_landmarks"]), integer(r["n_stays"]), integer(r["excluded_imminent_positive_landmarks"]), ci(r, "auroc"), ci(r, "auprc"), ci(r, "brier", 5)])
    return (weight_header, weight_data), (gap_header, gap_data)


def supplementary_8() -> tuple[list[str], list[list[str]]]:
    header = ["Database", "Calibration", "Threshold", "Suppression", "Landmarks", "Event stays", "Alert episodes", "True episodes", "False episodes", "Event-stay sensitivity", "Episode PPV", "False episodes / 100 patient-days"]
    data = []
    for r in rows("corrected_alert_suppression_metrics.csv"):
        data.append([label_dataset(r["dataset"]), r["calibration"], f(r["threshold"], 2), f(r["suppression_hours"], 0) + " h", integer(r["n_landmarks"]), integer(r["n_event_stays"]), integer(r["alert_episodes"]), integer(r["true_alert_episodes"]), integer(r["false_alert_episodes"]), f(r["event_stay_sensitivity"], 3), f(r["episode_ppv"], 3), f(r["false_alert_episodes_per_100_patient_days"], 2)])
    return header, data


def build() -> None:
    t1h, t1d = table_1()
    t2h, t2d = table_2()
    t3h, t3d = table_3()
    t4h, t4d = table_4()
    t5h, t5d = table_5()
    s1h, s1d = supplementary_1()
    s2h, s2d = supplementary_2()
    s3h, s3d = supplementary_3()
    s4h, s4d = supplementary_4()
    (s5ah, s5ad), (s5bh, s5bd) = supplementary_5()
    (s6ah, s6ad), (s6bh, s6bd), (s6ch, s6cd) = supplementary_6()
    (s7ah, s7ad), (s7bh, s7bd) = supplementary_7()
    s8h, s8d = supplementary_8()

    exports = [
        (TABLES / "Table_1_Cohort_characteristics.csv", t1h, t1d),
        (TABLES / "Table_2_Main_validation_results.csv", t2h, t2d),
        (TABLES / "Table_3_Calibration_and_site_heterogeneity.csv", t3h, t3d),
        (TABLES / "Table_4_Robustness_and_sensitivity.csv", t4h, t4d),
        (TABLES / "Table_5_Threshold_and_alert_policy.csv", t5h, t5d),
        (SUPP / "Supplementary_Table_1_Database_design.csv", s1h, s1d),
        (SUPP / "Supplementary_Table_2_Predictor_dictionary.csv", s2h, s2d),
        (SUPP / "Supplementary_Table_3_Subgroup_performance.csv", s3h, s3d),
        (SUPP / "Supplementary_Table_4_Recalibration_summary.csv", s4h, s4d),
         (SUPP / "Supplementary_Table_5_Horizon_and_estimands.csv", s5ah, s5ad),
         (SUPP / "Supplementary_Table_6_Clinical_comparators_and_reduced_models.csv", s6ah, s6ad),
         (SUPP / "Supplementary_Table_7_Training_and_action_gap.csv", s7ah, s7ad),
         (SUPP / "Supplementary_Table_8_Alert_suppression.csv", s8h, s8d),
         (SUPP / "Supplementary_Table_5A_Horizon_eligibility.csv", s5ah, s5ad),
         (SUPP / "Supplementary_Table_5B_Alternative_estimands.csv", s5bh, s5bd),
         (SUPP / "Supplementary_Table_6A_Clinical_comparators.csv", s6ah, s6ad),
         (SUPP / "Supplementary_Table_6B_Reduced_feature_models.csv", s6bh, s6bd),
         (SUPP / "Supplementary_Table_6C_Temperature_audit.csv", s6ch, s6cd),
         (SUPP / "Supplementary_Table_7A_Equal_total_stay_weight_training.csv", s7ah, s7ad),
         (SUPP / "Supplementary_Table_7B_One_hour_clinical_action_gap.csv", s7bh, s7bd),
     ]
    for path, header, data in exports:
        write_csv(path, header, data)
        write_markdown(path.with_suffix(".md"), path.stem.replace("_", " "), header, data)

    write_markdown(TABLES / "Table_1_Cohort_characteristics.md", "Table 1. Cohort characteristics", t1h, t1d, "Age is summarized at the stay level. Percentages are calculated from the audited cohort summary.")
    write_markdown(TABLES / "Table_2_Main_validation_results.md", "Table 2. Frozen-model validation", t2h, t2d, "Confidence intervals use 500 stay-clustered bootstrap replicates for AUROC, AUPRC and Brier score, and stay-clustered sandwich intervals for calibration statistics.")
    write_markdown(TABLES / "Table_3_Calibration_and_site_heterogeneity.md", "Table 3. Calibration and site-level heterogeneity", t3h, t3d, "CITL is calibration-in-the-large. The 73-hospital row is descriptive; the 208-hospital row uses hospital-clustered bootstrap inference.")
    write_markdown(TABLES / "Table_4_Robustness_and_sensitivity.md", "Table 4. Robustness and sensitivity analyses", t4h, t4d, "All sensitivity analyses were post hoc and were not used for model selection.")
    write_markdown(TABLES / "Table_5_Threshold_and_alert_policy.md", "Table 5. Threshold and alert-policy metrics", t5h, t5d, "Sens. = sensitivity; Spec. = specificity; PPV = positive predictive value. The 0.05 threshold and six-hour suppression policy are exploratory and do not establish prospective clinical benefit.")

    save_table_doc(TABLES / "Table_1_Cohort_characteristics.docx", "Table 1. Cohort characteristics and event prevalence", [(t1h, t1d, [1700, 950, 900, 1350, 950, 1250, 1200, 1050, 1010], "Age is shown as median (IQR). Event prevalence is reported at both the stay and landmark levels in the corresponding columns.")])
    save_table_doc(TABLES / "Table_2_Main_validation_results.docx", "Table 2. Main validation results for frozen models", [(t2h, t2d, [1350, 1250, 1200, 1050, 1400, 1400, 1400, 1300, 1310], "CITL is calibration-in-the-large with the prediction-logit slope fixed at 1. Negative values indicate average overprediction; positive values indicate average underprediction." )])
    save_table_doc(TABLES / "Table_3_Calibration_and_site_heterogeneity.docx", "Table 3. Calibration and site-level heterogeneity", [(t3h, t3d, [1550, 1900, 1050, 950, 1350, 1600, 1960], "CITL is calibration-in-the-large. The 73-hospital row is descriptive; the 208-hospital row uses hospital-clustered bootstrap inference." )])
    save_table_doc(TABLES / "Table_4_Robustness_and_sensitivity.docx", "Table 4. Robustness and sensitivity analyses", [(t4h, t4d, [2200, 1650, 1650, 1650, 2210], "All sensitivity analyses were post hoc and were not used for model selection." )])
    save_table_doc(TABLES / "Table_5_Threshold_and_alert_policy.docx", "Table 5. Threshold and alert-policy metrics", [(t5h, t5d, [1200, 1250, 950, 950, 950, 700, 1500, 1550, 900, 1410], "Sens. = sensitivity; Spec. = specificity; PPV = positive predictive value. The 0.05 threshold and six-hour suppression policy are exploratory and do not establish prospective clinical benefit." )])
    supp_sections = [
        (s1h, s1d, [850, 700, 1250, 1250, 1450, 1150, 1650, 1060], "Database versions and access terms are based on the audited data-availability record."),
        (s2h, s2d, [500, 1100, 1400, 800, 950, 700, 850, 1000, 1000, 1000, 1050, 1010], "The 42 predictors were fixed before external scoring. Values outside the accepted range were set to missing."),
        (s3h, s3d, [1100, 850, 900, 800, 700, 800, 1200, 1200, 1200, 750, 760], "Subgroup analyses were exploratory and were not used for model selection or formal fairness inference."),
        (s4h, s4d, [1250, 750, 900, 900, 1700, 1900, 1960], "Negative Brier changes indicate improvement after intercept recalibration. Each split was identifier-disjoint between calibration and evaluation subsets."),
        (s5ah, s5ad, [1100, 1500, 1600, 1550, 1100, 1250, 1260], "The primary binary estimand requires a complete six-hour horizon; incomplete landmarks were excluded."),
        (s5bh, s5bd, [1200, 2100, 1150, 900, 1500, 1500, 1510], "The hour-6 analysis uses one landmark per stay at ICU hour 6."),
         (s6ah, s6ad, [1200, 2400, 1300, 2200, 2260], "Clinical comparators were evaluated only where the required measurements were available."),
         (s6bh, s6bd, [1150, 2300, 1150, 1500, 1500, 1500, 1260], "Reduced-feature models used the fixed primary HGB hyperparameters and were fitted only in MIMIC-IV 2008-2019."),
         (s6ch, s6cd, [1200, 1200, 1500, 1300, 1500, 1500, 1300, 1300, 1300], "Temperature distributions were audited before the temperature-exclusion sensitivity analysis; percentages use all eligible landmarks as the denominator."),
         (s7ah, s7ad, [1250, 2200, 1200, 900, 1500, 1500, 1810], "Equal-total-stay weighting was a post hoc training sensitivity analysis."),
        (s7bh, s7bd, [1100, 2100, 1200, 900, 1500, 1500, 1500, 1560], "The one-hour action-gap target excludes initiation within one hour of the landmark and predicts initiation during hours (1, 6]."),
        (s8h, s8d, [1000, 1000, 850, 850, 1000, 950, 1050, 950, 950, 1250, 900, 1560], "This is one descriptive six-hour alert-suppression policy and does not establish prospective clinical benefit."),
    ]
    supp_titles = [
        "Supplementary Table 1. Database design and measurement context",
        "Supplementary Table 2. Predictor dictionary and source mappings",
        "Supplementary Table 3. Exploratory subgroup performance",
        "Supplementary Table 4. Repeated intercept recalibration",
        "Supplementary Table 5A. Complete-horizon eligibility audit",
        "Supplementary Table 5B. Alternative landmark estimands",
         "Supplementary Table 6A. Clinical comparators",
         "Supplementary Table 6B. Reduced-feature models",
         "Supplementary Table 6C. Temperature audit",
         "Supplementary Table 7A. Equal-total-stay-weight training",
        "Supplementary Table 7B. One-hour clinical-action gap",
        "Supplementary Table 8. Six-hour alert-suppression sensitivity",
    ]
    save_table_doc(
        SUPP / "Supplementary_Tables_1_to_8.docx",
        "Supplementary tables",
        supp_sections,
        section_titles=supp_titles,
        page_break_between=True,
    )

    # A compact index helps the author and editorial office map every table to its source data.
    index = [
        ["Table 1", "Cohort characteristics", "tables/Table_1_Cohort_characteristics.csv"],
        ["Table 2", "Frozen-model validation", "tables/Table_2_Main_validation_results.csv"],
        ["Table 3", "Calibration and site-level heterogeneity", "tables/Table_3_Calibration_and_site_heterogeneity.csv"],
        ["Table 4", "Robustness and sensitivity analyses", "tables/Table_4_Robustness_and_sensitivity.csv"],
        ["Table 5", "Threshold and alert-policy metrics", "tables/Table_5_Threshold_and_alert_policy.csv"],
        ["Supplementary Table 1", "Database design and measurement context", "supplementary/Supplementary_Table_1_Database_design.csv"],
        ["Supplementary Table 2", "42-predictor dictionary", "supplementary/Supplementary_Table_2_Predictor_dictionary.csv"],
        ["Supplementary Table 3", "Exploratory subgroup performance", "supplementary/Supplementary_Table_3_Subgroup_performance.csv"],
        ["Supplementary Table 4", "Repeated intercept recalibration", "supplementary/Supplementary_Table_4_Recalibration_summary.csv"],
         ["Supplementary Table 5", "Horizon and alternative estimands", "supplementary/Supplementary_Table_5_Horizon_and_estimands.csv"],
         ["Supplementary Table 6", "Clinical comparators and reduced models", "supplementary/Supplementary_Table_6_Clinical_comparators_and_reduced_models.csv"],
         ["Supplementary Table 7", "Training weighting and action gap", "supplementary/Supplementary_Table_7_Training_and_action_gap.csv"],
         ["Supplementary Table 8", "Alert suppression sensitivity", "supplementary/Supplementary_Table_8_Alert_suppression.csv"],
         ["Supplementary Table 5A", "Complete-horizon eligibility audit", "supplementary/Supplementary_Table_5A_Horizon_eligibility.csv"],
         ["Supplementary Table 5B", "Alternative landmark estimands", "supplementary/Supplementary_Table_5B_Alternative_estimands.csv"],
         ["Supplementary Table 6A", "Clinical comparators", "supplementary/Supplementary_Table_6A_Clinical_comparators.csv"],
         ["Supplementary Table 6B", "Reduced-feature models", "supplementary/Supplementary_Table_6B_Reduced_feature_models.csv"],
         ["Supplementary Table 6C", "Temperature audit", "supplementary/Supplementary_Table_6C_Temperature_audit.csv"],
         ["Supplementary Table 7A", "Equal-total-stay-weight training", "supplementary/Supplementary_Table_7A_Equal_total_stay_weight_training.csv"],
         ["Supplementary Table 7B", "One-hour clinical-action gap", "supplementary/Supplementary_Table_7B_One_hour_clinical_action_gap.csv"],
     ]
    write_csv(PKG / "table_source_index.csv", ["Item", "Description", "Editable source"], index)


if __name__ == "__main__":
    build()
